from docx import Document
from docx.shared import Cm
import os
import numpy as np

import sys

if len(sys.argv) != 2:
    print("Please use python3.7 Word.py \"yourgrahpname\"")
    sys.exit()

print("len(sys.argv): " + str(len(sys.argv)))

title = sys.argv[1]

document = Document()
document.add_heading(str(title), 0)

areaList = os.listdir(str(title))
print("areaList: " + str(areaList))

if ".DS_Store" in areaList:
    areaList.remove(".DS_Store")

areaList = np.sort(areaList)

for area in areaList:

    listFiles = os.listdir(str(title) + "/" + str(area))
    listFiles = np.sort(listFiles)

    listImage = []
    for file in listFiles:
        print("file: " + str(file))
        listImage.append(file)

    if ".DS_Store" in listImage:
        listImage.remove(".DS_Store")

    document.add_heading(area, level=1)
    for pngName in listImage:
        document.add_picture(str(title) + "/" + str(area) + "/" + str(pngName), width=Cm(20))
        document.add_page_break()

document.save(str(title) + '.docx')
